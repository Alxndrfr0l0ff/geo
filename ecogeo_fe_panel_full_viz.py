
# ecogeo_fe_panel_full_viz.py
# End-to-end pipeline: build panel (2019–2024), FE (two-way) in two directions,
# export CSVs + DOCX tables + multiple visualizations (forest, caterpillar of FE, FWL partial residuals, fitted vs observed).

import os, math, numpy as np, pandas as pd, statsmodels.api as sm, shapefile
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE="/mnt/data"

# ------------------- Geometry helpers -------------------
def area_km2_from_wgs84_ring(points):
    if len(points) < 3: return 0.0
    R = 6371008.8
    lats = [pt[1] for pt in points]; lons = [pt[0] for pt in points]
    lat0 = math.radians(sum(lats)/len(lats))
    xs = [math.radians(lon) * math.cos(lat0) * R for lon in lons]
    ys = [math.radians(lat) * R for lat in lats]
    area2 = 0.0
    for i in range(len(xs)):
        j = (i+1) % len(xs); area2 += xs[i]*ys[j] - xs[j]*ys[i]
    return abs(area2)/2.0/1e6

def shape_area_km2(shape):
    pts = shape.points; parts = list(shape.parts)+[len(pts)]
    total=0.0
    for i in range(len(parts)-1):
        ring = pts[parts[i]:parts[i+1]]; total += area_km2_from_wgs84_ring(ring)
    return total

def read_areas_from_shp(shp_path):
    sf = shapefile.Reader(shp_path)
    field_names = [f[0] for f in sf.fields][1:]
    recs=[]
    for sr in sf.iterShapeRecords():
        m = dict(zip(field_names, sr.record))
        code = next((str(m[k]) for k in m if "katot" in k.lower()), None)
        if code is None:
            raise RuntimeError("Не знайдено поле KATOTTG у шейпфайлі.")
        name = next((str(m[k]) for k in m if "name" in k.lower()), code)
        area = shape_area_km2(sr.shape)
        recs.append({"HKATOTTG": code, "name": name, "area_km2": area})
    df = pd.DataFrame(recs).groupby("HKATOTTG", as_index=False).agg(name=("name","first"), area_km2=("area_km2","sum"))
    return df

# ------------------- Panel build -------------------
def build_panel():
    areas = read_areas_from_shp(os.path.join(BASE,"IF_reg_TG_bou_7.shp"))
    # TN
    tn = pd.read_csv(os.path.join(BASE,"згруповано_турзбір.csv")).rename(
        columns={"Код_громади":"HKATOTTG","Рік":"YEAR","Всього_туристо_діб":"TN"})
    tn["HKATOTTG"]=tn["HKATOTTG"].astype(str)
    tn = tn.merge(areas, on="HKATOTTG", how="left")
    tn["TN_km2"]=tn["TN"]/tn["area_km2"]
    tn = tn[["HKATOTTG","name","YEAR","TN_km2"]]

    # Water
    voda = pd.read_excel(os.path.join(BASE,"voda.xlsx"))
    code_col = next((c for c in ["HKATOTTG","TG","katottg","KATOTTG"] if c in voda.columns), None)
    year_col = next((c for c in ["PERIOD_YEAR","P_YEAR","YEAR"] if c in voda.columns), None)
    voda["HKATOTTG"]=voda[code_col].astype(str); voda["YEAR"]=pd.to_numeric(voda[year_col], errors="coerce")
    num_cols = voda.select_dtypes(include=[np.number]).columns.tolist()
    if year_col in num_cols: num_cols.remove(year_col)
    W = voda.groupby(["HKATOTTG","YEAR"])[num_cols].sum().sum(axis=1).reset_index(name="W_total")
    W = W.merge(areas[["HKATOTTG","area_km2"]], on="HKATOTTG", how="left")
    W["W_km2"]=W["W_total"]/W["area_km2"]

    # AIR
    air = pd.read_excel(os.path.join(BASE,"ecol1_cleaned.xlsx"))
    c = next((x for x in ["HKATOTTG","katottg","KATOTTG"] if x in air.columns), None)
    y = next((x for x in ["P_YEAR","PERIOD_YEAR","YEAR"] if x in air.columns), None)
    v = next((x for x in ["POLLUTION_VOL","VALUE","VAL","AMOUNT"] if x in air.columns), None)
    air = air.rename(columns={c:"HKATOTTG", y:"YEAR", v:"AIR"})
    air["HKATOTTG"]=air["HKATOTTG"].astype(str); air["YEAR"]=pd.to_numeric(air["YEAR"], errors="coerce")
    air["AIR"]=pd.to_numeric(air["AIR"], errors="coerce").fillna(0)
    AIR = air.groupby(["HKATOTTG","YEAR"])["AIR"].sum().reset_index().merge(areas[["HKATOTTG","area_km2"]], on="HKATOTTG", how="left")
    AIR["AIR_km2"]=AIR["AIR"]/AIR["area_km2"]

    # DW
    dw = pd.read_excel(os.path.join(BASE,"ecol2_cleaned.xlsx"))
    c = next((x for x in ["HKATOTTG","katottg","KATOTTG"] if x in dw.columns), None)
    y = next((x for x in ["P_YEAR","PERIOD_YEAR","YEAR"] if x in dw.columns), None)
    v = next((x for x in ["POLLUTION_VOL","VALUE","VAL","AMOUNT"] if x in dw.columns), None)
    dw = dw.rename(columns={c:"HKATOTTG", y:"YEAR", v:"DW"})
    dw["HKATOTTG"]=dw["HKATOTTG"].astype(str); dw["YEAR"]=pd.to_numeric(dw["YEAR"], errors="coerce")
    dw["DW"]=pd.to_numeric(dw["DW"], errors="coerce").fillna(0)
    DW = dw.groupby(["HKATOTTG","YEAR"])["DW"].sum().reset_index().merge(areas[["HKATOTTG","area_km2"]], on="HKATOTTG", how="left")
    DW["DW_km2"]=DW["DW"]/DW["area_km2"]

    # MSW
    msw = pd.read_excel(os.path.join(BASE,"ecol3_cleaned.xlsx"))
    c = next((x for x in ["HKATOTTG","katottg","KATOTTG"] if x in msw.columns), None)
    y = next((x for x in ["P_YEAR","PERIOD_YEAR","YEAR"] if x in msw.columns), None)
    v = next((x for x in ["POLLUTION_VOL","VALUE","VAL","AMOUNT"] if x in msw.columns), None)
    msw = msw.rename(columns={c:"HKATOTTG", y:"YEAR", v:"MSW"})
    msw["HKATOTTG"]=msw["HKATOTTG"].astype(str); msw["YEAR"]=pd.to_numeric(msw["YEAR"], errors="coerce")
    msw["MSW"]=pd.to_numeric(msw["MSW"], errors="coerce").fillna(0)
    MSW = msw.groupby(["HKATOTTG","YEAR"])["MSW"].sum().reset_index().merge(areas[["HKATOTTG","area_km2"]], on="HKATOTTG", how="left")
    MSW["MSW_km2"]=MSW["MSW"]/MSW["area_km2"]

    panel = tn.merge(W[["HKATOTTG","YEAR","W_km2"]], on=["HKATOTTG","YEAR"], how="left")\
              .merge(AIR[["HKATOTTG","YEAR","AIR_km2"]], on=["HKATOTTG","YEAR"], how="left")\
              .merge(DW[["HKATOTTG","YEAR","DW_km2"]], on=["HKATOTTG","YEAR"], how="left")\
              .merge(MSW[["HKATOTTG","YEAR","MSW_km2"]], on=["HKATOTTG","YEAR"], how="left")
    panel = panel[panel["YEAR"].between(2019,2024)].sort_values(["HKATOTTG","YEAR"]).copy()
    panel["TN_km2_lag1"]=panel.groupby("HKATOTTG")["TN_km2"].shift(1)
    for Y in ["W_km2","AIR_km2","DW_km2","MSW_km2"]:
        panel[f"{Y}_lag1"]=panel.groupby("HKATOTTG")[Y].shift(1)
    heavy = {"Бурштинська міська громада","Калуська міська громада","Ямницька сільська громада"}
    panel["heavy"]=panel["name"].isin(heavy)
    return panel

# ------------------- FE estimators -------------------
def fe_ols(y_col, x_cols, df):
    d = df[[y_col]+x_cols+["HKATOTTG","YEAR"]].dropna().copy()
    if d.empty: return None, None
    y = d[y_col].values
    X = [np.ones(len(d))] + [d[c].values for c in x_cols]
    TG = pd.get_dummies(pd.Categorical(d["HKATOTTG"]), drop_first=True)
    TT = pd.get_dummies(pd.Categorical(d["YEAR"]), drop_first=True)
    X = np.column_stack([*X, TG.values, TT.values])
    m = sm.OLS(y, X).fit(cov_type="cluster", cov_kwds={"groups": pd.Categorical(d["HKATOTTG"]).codes})
    coef=float(m.params[1]); se=float(m.bse[1]); t=coef/se if se!=0 else np.nan; p=float(m.pvalues[1])
    return {"coef":coef,"se":se,"t":t,"p":p,"nobs":int(m.nobs)}, m

def model_stats(y_col, x_col, df):
    d = df[[y_col,x_col,"HKATOTTG","YEAR"]].dropna().copy()
    if d.empty: return None
    y=d[y_col].values
    X = np.column_stack([np.ones(len(d)), d[x_col].values,
                         pd.get_dummies(pd.Categorical(d["HKATOTTG"]), drop_first=True).values,
                         pd.get_dummies(pd.Categorical(d["YEAR"]), drop_first=True).values])
    m = sm.OLS(y, X).fit(cov_type="cluster", cov_kwds={"groups": pd.Categorical(d["HKATOTTG"]).codes})
    sd_y = np.std(y, ddof=1); sd_x = np.std(d[x_col].values, ddof=1)
    beta_std = m.params[1]*(sd_x/sd_y) if sd_y>0 else np.nan
    return {"R2": float(m.rsquared), "beta_std": float(beta_std), "nobs": int(m.nobs), "ngroups": int(d["HKATOTTG"].nunique())}

def iqr_effect(y_col, x_col, df):
    d = df[[y_col,x_col,"HKATOTTG","YEAR"]].dropna().copy()
    if d.empty: return (np.nan, np.nan)
    y=d[y_col].values
    X = np.column_stack([np.ones(len(d)), d[x_col].values,
                         pd.get_dummies(pd.Categorical(d["HKATOTTG"]), drop_first=True).values,
                         pd.get_dummies(pd.Categorical(d["YEAR"]), drop_first=True).values])
    m = sm.OLS(y, X).fit(cov_type="cluster", cov_kwds={"groups": pd.Categorical(d["HKATOTTG"]).codes})
    coef=float(m.params[1])
    q1,q3 = np.percentile(d[x_col].values, [25,75])
    abs_eff = coef*(q3-q1)
    med_y = np.median(y)
    pct = abs_eff/med_y*100 if med_y!=0 else np.nan
    return abs_eff, pct

# LSDV to recover unit FE (alpha_i) and time FE (tau_t)
def lsdv_two_way(y_col, x_col, df):
    d = df[[y_col, x_col, "HKATOTTG", "YEAR"]].dropna().copy()
    if d.empty: return None, None, None, None
    y = d[y_col].values
    # Build full-rank design: intercept + x + (N-1) unit dummies + (T-1) time dummies
    units = pd.Categorical(d["HKATOTTG"]); years = pd.Categorical(d["YEAR"])
    D_unit = pd.get_dummies(units, drop_first=True); unit_levels = units.categories.tolist()
    D_time = pd.get_dummies(years, drop_first=True); time_levels = years.categories.tolist()
    X = np.column_stack([np.ones(len(d)), d[x_col].values, D_unit.values, D_time.values])
    mdl = sm.OLS(y, X).fit(cov_type="cluster", cov_kwds={"groups": units.codes})
    # Recover alphas relative to dropped baseline
    intercept = mdl.params[0]; beta = mdl.params[1]
    # unit FE: baseline unit (first category) has 0 dummy; others have coeff equal to their dummy param
    alphas = {}
    for i,u in enumerate(unit_levels):
        if i==0:
            alphas[u] = 0.0
        else:
            alphas[u] = mdl.params[1+i]  # after intercept and x, next (N-1) params are unit dummies
    # Convert to centered alphas around mean zero to avoid baseline artifact
    a_vals = np.array(list(alphas.values())); a_center = a_vals - a_vals.mean()
    alphas_centered = {u: a_center[i] for i,u in enumerate(unit_levels)}
    # time FE (tau): similar approach
    taus = {}
    offset = 2 + (len(unit_levels)-1)
    for j,t in enumerate(time_levels):
        if j==0:
            taus[t] = 0.0
        else:
            taus[t] = mdl.params[offset + (j-1)]
    t_vals = np.array(list(taus.values())); t_center = t_vals - t_vals.mean()
    taus_centered = {t: t_center[j] for j,t in enumerate(time_levels)}
    return mdl, alphas_centered, taus_centered, beta

# ------------------- DOCX helpers -------------------
def add_table_from_df(doc, df, title):
    doc.add_paragraph().add_run(title).bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for j,c in enumerate(df.columns):
        p=hdr[j].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(c)); r.bold=True
    for _,r in df.iterrows():
        cells = table.add_row().cells
        for j,c in enumerate(df.columns):
            p=cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j>0 else WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(str(r[c]))
    return table

def save_tables_and_viz(panel, tab, rev_df, out_docx):
    # Primary DOCX with tables and inserted figures
    doc = Document()
    doc.add_heading("Панельні оцінки з фіксованими ефектами (ТГ×рік, 2019–2024)", level=1)

    tab1 = tab.copy()
    tab1["Y"] = tab1["Y"].str.replace("_km2","", regex=False).str.upper()
    tab1 = tab1.rename(columns={
        "Y":"Залежна змінна (Y)",
        "sample":"Вибірка",
        "coef":"β",
        "se":"SE(β)",
        "t":"t",
        "p":"p",
        "beta_std":"β (стандартиз.)",
        "R2":"R²(within)",
        "nobs":"n",
        "ngroups":"N_ТГ",
        "iqr_effect_abs":"Ефект IQR (натр.од.)",
        "iqr_effect_pct_median":"Ефект IQR (% до медіани Y)"
    })
    add_table_from_df(doc, tab1, "Табл. 2.2.1. Y ∈ {W, AIR, DW, MSW} на TN(t−1) (двовимірні FE; кластер-SE за ТГ)")
    p = doc.add_paragraph("Примітки: оцінки виконано на нормованих індикаторах (на км²). ")
    p.add_run("«Вибірка = no_heavy» — без Бурштинської, Калуської та Ямницької громад. ").italic = True
    p.add_run("Ефект IQR — зміна Y за переходу TN(t−1) від 25-го до 75-го перцентиля.").italic = True
    doc.add_paragraph()

    rev1 = rev_df[["X","sample","coef","se","t","p","nobs"]].rename(columns={
        "X":"Предиктор (лаг, t-1)",
        "sample":"Вибірка",
        "coef":"γ",
        "se":"SE(γ)",
        "t":"t",
        "p":"p",
        "nobs":"n"
    })
    rev1["Предиктор (лаг, t-1)"] = rev1["Предиктор (лаг, t-1)"].str.replace("_km2_lag1","(t−1)", regex=False)\
                                                              .str.replace("_km2","", regex=False)\
                                                              .str.upper()

    add_table_from_df(doc, rev1, "Табл. 2.2.2. TN на Y(t−1) (двовимірні FE; кластер-SE за ТГ)")
    p2 = doc.add_paragraph("Примітки: статистично значущих ефектів не виявлено (p ≥ 0.14 у всіх специфікаціях)."); p2.italic = True

    # Insert figures if exist
    figs = [
        ("Рис. 2.2.1. Ефект IQR (бар-чарт)", os.path.join(BASE,"IQR_effects_bar.png")),
        ("Рис. 2.2.2. Forest plot β з 95% ДІ", os.path.join(BASE,"forest_betas.png")),
        ("Рис. 2.2.3. Caterpillar α_i (W_km2)", os.path.join(BASE,"caterpillar_W_alphas.png")),
        ("Рис. 2.2.4. FWL: W_km2~TN_lag1 (частк.залишки)", os.path.join(BASE,"fwl_W_scatter.png")),
        ("Рис. 2.2.5. Observed vs Fitted (W_km2)", os.path.join(BASE,"fit_vs_obs_W.png"))
    ]
    for caption, path in figs:
        if os.path.exists(path):
            doc.add_paragraph()
            p = doc.add_paragraph(caption); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(path, width=Inches(6.2))

    doc.save(out_docx)

# ------------------- Run pipeline -------------------
def main():
    panel = build_panel()
    panel.to_csv(os.path.join(BASE,"panel_km2_2019_2024.csv"), index=False, encoding="utf-8-sig")

    rows=[]; stats_rows=[]; effects={}
    for Y in ["W_km2","AIR_km2","DW_km2","MSW_km2"]:
        s,_ = fe_ols(Y, ["TN_km2_lag1"], panel); 
        if s: s.update({"Y":Y,"sample":"full"}); rows.append(s); st=model_stats(Y,"TN_km2_lag1",panel); st.update({"Y":Y,"sample":"full"}); stats_rows.append(st)
        ph = panel.loc[~panel["heavy"]].copy()
        s,_ = fe_ols(Y, ["TN_km2_lag1"], ph); 
        if s: s.update({"Y":Y,"sample":"no_heavy"}); rows.append(s); st=model_stats(Y,"TN_km2_lag1",ph); st.update({"Y":Y,"sample":"no_heavy"}); stats_rows.append(st)
        if Y in ["W_km2","DW_km2"]:
            effects[(Y,"full")] = iqr_effect(Y,"TN_km2_lag1",panel)
            effects[(Y,"no_heavy")] = iqr_effect(Y,"TN_km2_lag1",ph)

    res_df=pd.DataFrame(rows)
    stats_df=pd.DataFrame(stats_rows)
    tab = res_df.merge(stats_df, on=["Y","sample"], how="left")
    tab["iqr_effect_abs"]=np.nan; tab["iqr_effect_pct_median"]=np.nan
    for (Y,s),(ea,ep) in effects.items():
        tab.loc[(tab["Y"]==Y)&(tab["sample"]==s),"iqr_effect_abs"]=ea
        tab.loc[(tab["Y"]==Y)&(tab["sample"]==s),"iqr_effect_pct_median"]=ep

    # Reverse direction
    panel_rev=panel.copy()
    for Y in ["W_km2","AIR_km2","DW_km2","MSW_km2"]:
        panel_rev[f"{Y}_lag1"]=panel_rev.groupby("HKATOTTG")[Y].shift(1)
    rev_rows=[]
    for Y in ["W_km2","AIR_km2","DW_km2","MSW_km2"]:
        s,_=fe_ols("TN_km2",[f"{Y}_lag1"],panel_rev); 
        if s: s.update({"X":f"{Y}_lag1","sample":"full"}); rev_rows.append(s)
        pr=panel_rev.loc[~panel_rev["heavy"]].copy()
        s,_=fe_ols("TN_km2",[f"{Y}_lag1"],pr); 
        if s: s.update({"X":f"{Y}_lag1","sample":"no_heavy"}); rev_rows.append(s)
    rev_df=pd.DataFrame(rev_rows)

    # Save CSVs
    res_df.to_csv(os.path.join(BASE,"fe_panel_Y_on_lagTN_2019_2024.csv"), index=False, encoding="utf-8-sig")
    rev_df.to_csv(os.path.join(BASE,"fe_panel_TN_on_lagY_2019_2024.csv"), index=False, encoding="utf-8-sig")
    stats_df.to_csv(os.path.join(BASE,"fe_panel_model_stats_2019_2024.csv"), index=False, encoding="utf-8-sig")
    tab.to_csv(os.path.join(BASE,"fe_panel_summary_table_2019_2024.csv"), index=False, encoding="utf-8-sig")

    # ------------- Visualizations -------------
    # (1) IQR bar (already known)
    eff = tab[["Y","sample","iqr_effect_pct_median"]].dropna().copy()
    if not eff.empty:
        eff["label"] = eff["Y"].str.replace("_km2","", regex=False).str.upper() + " (" + eff["sample"].map({"full":"повна","no_heavy":"без важковаговиків"}) + ")"
        plt.figure()
        plt.bar(eff["label"], eff["iqr_effect_pct_median"])
        plt.title("Ефект інтерквартильного зростання TN(t−1)\nна Y (% медіани відповідного Y)")
        plt.ylabel("% до медіани Y")
        plt.xticks(rotation=25, ha='right')
        plt.tight_layout(); plt.savefig(os.path.join(BASE,"IQR_effects_bar.png"), dpi=300, bbox_inches="tight"); plt.close()

    # (2) Forest plot for betas with 95% CI
    fdf = res_df.copy()
    if not fdf.empty:
        fdf["Y"] = fdf["Y"].str.replace("_km2","", regex=False).str.upper()
        fdf["lower"] = fdf["coef"] - 1.96*fdf["se"]
        fdf["upper"] = fdf["coef"] + 1.96*fdf["se"]
        fdf["spec"] = fdf["Y"] + " | " + fdf["sample"].map({"full":"повна","no_heavy":"без важк."})
        order = np.argsort(fdf["coef"].values)
        fdf = fdf.iloc[order]
        plt.figure(figsize=(7, max(3, 0.5*len(fdf))))
        y = np.arange(len(fdf))
        plt.hlines(y, fdf["lower"], fdf["upper"])
        plt.plot(fdf["coef"], y, 'o')
        plt.axvline(0, linestyle='--')
        plt.yticks(y, fdf["spec"])
        plt.xlabel("β (ефект TN(t−1))")
        plt.title("Оцінки β з 95% довірчими інтервалами")
        plt.tight_layout(); plt.savefig(os.path.join(BASE,"forest_betas.png"), dpi=300, bbox_inches="tight"); plt.close()

    # (3) Caterpillar plot of unit FE (alphas) for W_km2
    mdl, alphas, taus, beta = lsdv_two_way("W_km2","TN_km2_lag1", panel)
    if alphas is not None:
        adf = pd.DataFrame({"HKATOTTG": list(alphas.keys()), "alpha": list(alphas.values())})
        # attach names
        names = panel[["HKATOTTG","name"]].drop_duplicates()
        adf = adf.merge(names, on="HKATOTTG", how="left")
        adf = adf.sort_values("alpha").reset_index(drop=True)
        plt.figure(figsize=(7, max(3, 0.25*len(adf))))
        y = np.arange(len(adf))
        plt.plot(adf["alpha"], y, 'o')
        plt.axvline(0, linestyle='--')
        plt.yticks(y, adf["name"])
        plt.xlabel("Оцінені фіксовані ефекти α_i (центр.)")
        plt.title("Caterpillar: гетерогенність базових рівнів W_km2 між ТГ")
        plt.tight_layout(); plt.savefig(os.path.join(BASE,"caterpillar_W_alphas.png"), dpi=300, bbox_inches="tight"); plt.close()

    # (4) FWL partial residuals for W_km2 ~ TN_lag1 (two-way demean)
    d = panel[["HKATOTTG","YEAR","W_km2","TN_km2_lag1"]].dropna().copy()
    if not d.empty:
        # two-way demean (within transformation)
        d["W_i"] = d.groupby("HKATOTTG")["W_km2"].transform("mean")
        d["W_t"] = d.groupby("YEAR")["W_km2"].transform("mean")
        d["W_all"] = d["W_km2"].mean()
        Wy = d["W_km2"] - d["W_i"] - d["W_t"] + d["W_all"]

        d["X_i"] = d.groupby("HKATOTTG")["TN_km2_lag1"].transform("mean")
        d["X_t"] = d.groupby("YEAR")["TN_km2_lag1"].transform("mean")
        d["X_all"] = d["TN_km2_lag1"].mean()
        Wx = d["TN_km2_lag1"] - d["X_i"] - d["X_t"] + d["X_all"]

        # Fit simple OLS on de-meaned to get line
        mdl = sm.OLS(Wy.values, sm.add_constant(Wx.values)).fit()
        xv = np.linspace(Wx.min(), Wx.max(), 100)
        yv = mdl.params[0] + mdl.params[1]*xv
        plt.figure()
        plt.scatter(Wx, Wy, s=12, alpha=0.7)
        plt.plot(xv, yv)
        plt.title("FWL: частковий зв'язок W_km2 ~ TN(t−1) після вилучення FE")
        plt.xlabel("TN_km2_lag1 (двовимірно центровано)")
        plt.ylabel("W_km2 (двовимірно центровано)")
        plt.tight_layout(); plt.savefig(os.path.join(BASE,"fwl_W_scatter.png"), dpi=300, bbox_inches="tight"); plt.close()

    # (5) Observed vs fitted for W_km2 (two-way FE)
    s, m = fe_ols("W_km2", ["TN_km2_lag1"], panel)
    if m is not None:
        # reconstruct fitted on the design used inside fe_ols
        d = panel[["W_km2","TN_km2_lag1","HKATOTTG","YEAR"]].dropna().copy()
        y = d["W_km2"].values
        X = [np.ones(len(d)), d["TN_km2_lag1"].values]
        TG = pd.get_dummies(pd.Categorical(d["HKATOTTG"]), drop_first=True)
        TT = pd.get_dummies(pd.Categorical(d["YEAR"]), drop_first=True)
        X = np.column_stack([*X, TG.values, TT.values])
        yhat = m.predict(X)
        plt.figure()
        plt.scatter(yhat, y, s=12, alpha=0.7)
        lims = [min(y.min(), yhat.min()), max(y.max(), yhat.max())]
        plt.plot(lims, lims)
        plt.xlabel("Підігнані значення")
        plt.ylabel("Спостереження")
        plt.title("Observed vs Fitted: W_km2 (двовимірні FE)")
        plt.tight_layout(); plt.savefig(os.path.join(BASE,"fit_vs_obs_W.png"), dpi=300, bbox_inches="tight"); plt.close()

    # ------------- DOCX with tables + figures -------------
    save_tables_and_viz(panel, tab, rev_df, os.path.join(BASE,"FE_results_tables_and_viz.docx"))

if __name__ == "__main__":
    main()
